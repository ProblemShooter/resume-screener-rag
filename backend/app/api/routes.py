import logging
from fastapi import APIRouter, HTTPException, Query, File, UploadFile, Form
from typing import List
from backend.app.api.schemas import (
    CandidateResponse, ScreenRequest, ScreenResponse, ScreenResult,
    ChatRequest, ChatResponse, ExplainResponse, ChunkScore, MatchDetails,
    CandidateAddRequest
)
from backend.app.services.db_service import db_service
from backend.app.services.embed_service import embed_service
from backend.app.services.llm_service import llm_service
from backend.app.config import settings

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extracts text from an uploaded file (PDF, DOCX, or TXT).
    """
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    elif ext == "docx":
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        text = ""
        for p in doc.paragraphs:
            if p.text.strip():
                text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        return text
    else:
        # Default to raw text decoding
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_content.decode("latin-1")
            except Exception as e:
                raise ValueError(f"Unsupported file format or encoding: {str(e)}")

logger = logging.getLogger(__name__)
router = APIRouter()

# In-memory cache for evaluation details to avoid re-calling Groq unnecessarily
EVALUATION_CACHE = {}

@router.get("/candidates", response_model=List[CandidateResponse])
def get_candidates():
    """
    Returns the list of all available candidates in the vector database.
    """
    try:
        candidates = db_service.get_all_candidates()
        return candidates
    except Exception as e:
        logger.error(f"Error listing candidates: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

@router.post("/screen", response_model=ScreenResponse)
def screen_candidates(req: ScreenRequest):
    """
    Two-stage screening pipeline:
    1. ChromaDB vector search (retrieve candidate chunks)
    2. Cross-Encoder re-ranking
    3. Groq LLM evaluation on top ranked candidates
    """
    if not req.job_description.strip():
        raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
    try:
        # --- STAGE 1: Retrieve candidate chunks ---
        # Fetch matching chunks. The db_service search returns raw cosine similarity scores.
        matched_chunks = db_service.search(
            query_text=req.job_description,
            category=req.category,
            limit=25 # Retrieve plenty of chunks
        )
        
        if not matched_chunks:
            return ScreenResponse(job_description=req.job_description, results=[], total_screened=0)
            
        # Group chunk matches by candidate
        candidate_groups = {}
        for chunk in matched_chunks:
            cid = chunk["candidate_id"]
            if cid not in candidate_groups:
                candidate_groups[cid] = {
                    "candidate_id": cid,
                    "candidate_name": chunk["candidate_name"],
                    "email": chunk["email"],
                    "category": chunk["category"],
                    "best_similarity": chunk["similarity"],
                    "chunks": []
                }
            candidate_groups[cid]["chunks"].append(chunk)
            
        # For each candidate, fetch their full resume text for stage 2 and LLM analysis
        candidates_to_rank = []
        for cid, cand_data in candidate_groups.items():
            full_chunks = db_service.get_candidate_chunks(cid)
            full_text = "\n".join([c["text"] for c in full_chunks])
            cand_data["full_text"] = full_text
            candidates_to_rank.append(cand_data)
            
        # --- STAGE 2: Cross-Encoder Reranking ---
        documents = [c["full_text"] for c in candidates_to_rank]
        cids = [c["candidate_id"] for c in candidates_to_rank]
        
        reranked_results = embed_service.rerank(
            query=req.job_description,
            documents=documents,
            candidate_ids=cids
        )
        
        # Build dictionary of scores
        rerank_scores_map = {r["candidate_id"]: r["rerank_score"] for r in reranked_results}
        
        # Sort candidates by rerank score descending
        candidates_to_rank = sorted(
            candidates_to_rank,
            key=lambda x: rerank_scores_map.get(x["candidate_id"], 0.0),
            reverse=True
        )
        
        # --- STAGE 3: Groq LLM Evaluation ---
        # Evaluate top 8 candidates for speed, provide estimates for lower ones
        candidates_to_evaluate = []
        evaluation_results = {}
        
        for idx, cand in enumerate(candidates_to_rank):
            cid = cand["candidate_id"]
            cache_key = f"{cid}_{hash(req.job_description)}"
            if cache_key in EVALUATION_CACHE:
                evaluation_results[cid] = EVALUATION_CACHE[cache_key]
            elif idx < 8:
                candidates_to_evaluate.append(cand)
            else:
                rerank_score = rerank_scores_map.get(cid, 0.0)
                evaluation_results[cid] = {
                    "overall_match_percentage": int(rerank_score),
                    "skills_match_percentage": int(rerank_score * 0.95),
                    "experience_match_percentage": int(rerank_score * 0.9),
                    "education_match_percentage": 70,
                    "project_match_percentage": 60,
                    "strengths": ["Matches standard keywords", "Relevant background"],
                    "weaknesses": ["Further details require manual evaluation"],
                    "missing_skills": [],
                    "verdict_summary": "Passed preliminary vector ranking. Screen further if top candidates are unsuitable."
                }
                
        # Run LLM evaluations in parallel for candidates not cached
        if candidates_to_evaluate:
            from concurrent.futures import ThreadPoolExecutor
            def run_eval(cand):
                cid = cand["candidate_id"]
                cache_key = f"{cid}_{hash(req.job_description)}"
                try:
                    eval_details = llm_service.evaluate_candidate_vs_jd(
                        candidate_name=cand["candidate_name"],
                        resume_text=cand["full_text"],
                        job_description=req.job_description
                    )
                    EVALUATION_CACHE[cache_key] = eval_details
                    return cid, eval_details
                except Exception as ex:
                    logger.error(f"Error in thread evaluating candidate {cand['candidate_name']}: {str(ex)}")
                    fallback = {
                        "overall_match_percentage": 50,
                        "skills_match_percentage": 50,
                        "experience_match_percentage": 50,
                        "education_match_percentage": 50,
                        "project_match_percentage": 50,
                        "strengths": ["Strong engineering interest"],
                        "weaknesses": ["Requires manual evaluation due to API timeout"],
                        "missing_skills": [],
                        "verdict_summary": f"Could not perform complete AI evaluation. Error: {str(ex)}"
                    }
                    EVALUATION_CACHE[cache_key] = fallback
                    return cid, fallback

            max_workers = min(8, len(candidates_to_evaluate))
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                thread_results = list(executor.map(run_eval, candidates_to_evaluate))
                for cid, eval_details in thread_results:
                    evaluation_results[cid] = eval_details

        # Build list of ScreenResult in sorted order
        results = []
        for idx, cand in enumerate(candidates_to_rank):
            cid = cand["candidate_id"]
            rerank_score = rerank_scores_map.get(cid, 0.0)
            initial_score = cand["best_similarity"]
            rank = idx + 1
            eval_details = evaluation_results.get(cid)
            
            results.append(ScreenResult(
                candidate_id=cid,
                candidate_name=cand["candidate_name"],
                email=cand["email"],
                category=cand["category"],
                initial_score=initial_score,
                rerank_score=rerank_score,
                rank=rank,
                evaluation=MatchDetails(**eval_details)
            ))
            
        return ScreenResponse(
            job_description=req.job_description,
            results=results,
            total_screened=len(results)
        )
        
    except Exception as e:
        logger.error(f"Error screening candidates: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Screening error: {str(e)}")

@router.post("/chat", response_model=ChatResponse)
def chat_candidate(req: ChatRequest):
    """
    Chat directly with a candidate's resume (Single Candidate RAG).
    """
    try:
        # Retrieve all chunks belonging to the candidate
        chunks = db_service.get_candidate_chunks(req.candidate_id)
        if not chunks:
            raise HTTPException(status_code=404, detail="Candidate not found")
            
        candidate_name = db_service.get_candidate_chunks(req.candidate_id)[0].get("candidate_name", "Candidate")
        
        # Send query to llm_service
        history_list = [{"role": msg.role, "content": msg.content} for msg in req.history]
        history_list.append({"role": "user", "content": req.message})
        
        answer = llm_service.chat_with_resume(
            candidate_name=candidate_name,
            resume_chunks=chunks,
            message_history=history_list[:-1] # History excluding the latest message which is sent separately
        )
        
        return ChatResponse(response=answer)
    except Exception as e:
        logger.error(f"Error chatting with candidate: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/candidates/{id}/questions", response_model=List[str])
def get_interview_questions(id: str, jd: str = Query(..., description="Job Description context")):
    """
    Generates customized interview questions based on candidate skill gaps relative to the JD.
    """
    try:
        chunks = db_service.get_candidate_chunks(id)
        if not chunks:
            raise HTTPException(status_code=404, detail="Candidate not found")
            
        candidate_name = chunks[0].get("candidate_name", "Candidate")
        full_text = "\n".join([c["text"] for c in chunks])
        
        questions = llm_service.generate_custom_interview_questions(
            candidate_name=candidate_name,
            resume_text=full_text,
            job_description=jd
        )
        return questions
    except Exception as e:
        logger.error(f"Error generating interview questions: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/candidates/{id}/explain", response_model=ExplainResponse)
def explain_candidate_score(id: str, jd: str = Query(..., description="Job Description context")):
    """
    Explains the score by returning the exact chunk-level match scores and evaluation details.
    """
    try:
        # 1. Retrieve all candidate chunks
        chunks = db_service.get_candidate_chunks(id)
        if not chunks:
            raise HTTPException(status_code=404, detail="Candidate not found")
            
        candidate_name = chunks[0].get("candidate_name", "Candidate")
        full_text = "\n".join([c["text"] for c in chunks])
        
        # 2. Score individual chunks using Cross-Encoder against the JD
        chunk_texts = [c["text"] for c in chunks]
        pairs = [[jd, text] for text in chunk_texts]
        scores = embed_service.rerank_model.predict(pairs)
        
        matching_chunks = []
        for chunk, score in zip(chunks, scores):
            # Sigmoid normalization
            import math
            norm_score = round((1 / (1 + math.exp(-score))) * 100, 2)
            
            matching_chunks.append(ChunkScore(
                section_name=chunk["section_name"],
                text=chunk["text"],
                similarity=norm_score
            ))
            
        # Sort chunks by matching similarity score descending
        matching_chunks = sorted(matching_chunks, key=lambda x: x.similarity, reverse=True)
        
        # 3. Pull or run evaluation
        cache_key = f"{id}_{hash(jd)}"
        eval_details = EVALUATION_CACHE.get(cache_key)
        if not eval_details:
            eval_details = llm_service.evaluate_candidate_vs_jd(
                candidate_name=candidate_name,
                resume_text=full_text,
                job_description=jd
            )
            EVALUATION_CACHE[cache_key] = eval_details
            
        return ExplainResponse(
            candidate_id=id,
            candidate_name=candidate_name,
            matching_chunks=matching_chunks,
            evaluation=MatchDetails(**eval_details)
        )
    except Exception as e:
        logger.error(f"Error explaining candidate score: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/candidates", response_model=dict)
def add_candidate_text(req: CandidateAddRequest):
    """
    Directly index a candidate's resume from raw text.
    """
    if not req.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text cannot be empty")
    try:
        result = db_service.add_new_candidate(req.resume_text, req.category)
        return result
    except Exception as e:
        logger.error(f"Error adding candidate: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/candidates/upload", response_model=dict)
async def upload_candidate_resume(
    file: UploadFile = File(..., description="Upload PDF, DOCX, or TXT resume file"),
    category: str = Form("General", description="Category for the candidate")
):
    """
    Upload and incrementally index a candidate's resume file (PDF, DOCX, or TXT).
    """
    try:
        content = await file.read()
        resume_text = extract_text_from_file(content, file.filename)
        
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Extracted resume text is empty")
            
        result = db_service.add_new_candidate(resume_text, category, is_uploaded=True)
        return result
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        logger.error(f"Error uploading resume: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"File processing error: {str(e)}")

@router.get("/stats", response_model=dict)
def get_database_stats():
    """
    Returns statistics about indexed candidate resumes vs total available resumes in the CSV.
    """
    try:
        indexed = len(db_service.get_all_candidates())
        total = db_service.get_csv_total_count()
        return {
            "indexed_resumes": indexed,
            "total_resumes_available": total,
            "dev_max_resumes": settings.DEV_MAX_RESUMES
        }
    except Exception as e:
        logger.error(f"Error fetching stats: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/settings/repopulate", response_model=dict)
def repopulate_database(max_resumes: int = Query(25, ge=1, le=1000, description="Number of resumes to index")):
    """
    Clears collection and re-populates it with the specified limit of resumes.
    """
    try:
        result = db_service.clear_and_repopulate_database(max_rows=max_resumes)
        return result
    except Exception as e:
        logger.error(f"Error repopulating database: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Repopulation failed: {str(e)}")
